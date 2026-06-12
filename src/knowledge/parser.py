"""
文档解析器模块

支持从多种文件格式中提取纯文本内容：
- TXT  : 直接 UTF-8 解码
- PDF  : pypdf 逐页提取
- DOCX : python-docx 段落提取
- XLSX : openpyxl 单元格遍历

Usage:
    from src.knowledge.parser import parse_document
    text = parse_document(file_bytes, "report.pdf")
"""

import io
import logging
from typing import Optional

logger = logging.getLogger("DocumentParser")
logger.setLevel(logging.DEBUG)
if not logger.handlers:
    _h = logging.StreamHandler()
    _h.setFormatter(logging.Formatter(
        "%(asctime)s [%(levelname)-7s] [%(name)s] %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    ))
    logger.addHandler(_h)

# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".xlsx"}


def _get_extension(filename: str) -> str:
    """提取小写文件扩展名，如 '.pdf'。"""
    idx = filename.rfind(".")
    if idx == -1:
        return ""
    return filename[idx:].lower()


def parse_document(file_bytes: bytes, filename: str) -> str:
    """
    根据文件扩展名自动选择解析器，提取纯文本。

    Args:
        file_bytes: 文件的原始字节内容
        filename: 文件名（用于判断扩展名）

    Returns:
        str: 提取的纯文本

    Raises:
        ValueError: 不支持的文件格式
    """
    ext = _get_extension(filename)
    logger.info("解析文档: filename=%s, ext=%s, size=%d bytes", filename, ext, len(file_bytes))

    if ext == ".txt":
        return _parse_txt(file_bytes)
    elif ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext == ".docx":
        return _parse_docx(file_bytes)
    elif ext == ".xlsx":
        return _parse_excel(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式: {ext}（支持: {', '.join(SUPPORTED_EXTENSIONS)}）")


# ================================================================
# 各格式解析器
# ================================================================

def _parse_txt(file_bytes: bytes) -> str:
    """TXT 文件：尝试 UTF-8 / GBK 解码。"""
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError("无法解码 TXT 文件，请确认文件编码为 UTF-8 或 GBK")


def _parse_pdf(file_bytes: bytes) -> str:
    """
    PDF 文件：使用 pypdf 逐页提取文本。

    Note:
        - 扫描版 PDF（纯图片）无法提取文字，返回空字符串
        - 提取后可能残留较多空白字符，需配合 cleaner 清洗
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text.strip())
        except Exception as e:
            logger.warning("PDF 第 %d 页提取失败: %s", i + 1, e)
            continue

    text = "\n\n".join(pages_text)
    logger.info("PDF 解析完成: 共 %d 页, 提取文本 %d 字符", len(reader.pages), len(text))
    return text


def _parse_docx(file_bytes: bytes) -> str:
    """
    DOCX 文件：使用 python-docx 提取段落和表格文本。

    提取顺序：段落文本 → 表格单元格文本
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # 段落文本
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # 表格文本
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts)
    logger.info("DOCX 解析完成: %d 段落, %d 表格, 提取文本 %d 字符",
                len(doc.paragraphs), len(doc.tables), len(text))
    return text


def _parse_excel(file_bytes: bytes) -> str:
    """
    XLSX 文件：使用 openpyxl 遍历所有 sheet 的单元格。

    按行输出，单元格用 tab 分隔，保留表格结构感。
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sheets_text = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines = [f"[Sheet: {sheet_name}]"]

        for row in ws.iter_rows(values_only=True):
            # 过滤全空行
            row_values = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(row_values):
                lines.append("\t".join(row_values))

        sheets_text.append("\n".join(lines))

    wb.close()
    text = "\n\n".join(sheets_text)
    logger.info("XLSX 解析完成: %d sheet, 提取文本 %d 字符", len(wb.sheetnames), len(text))
    return text
